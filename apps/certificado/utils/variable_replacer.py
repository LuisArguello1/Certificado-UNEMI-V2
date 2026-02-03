"""
Utilidad para el reemplazo de variables en plantillas Word (.docx).

Este módulo maneja la sustitución de marcadores {{VARIABLE}} preservando
el formato original del documento. También aplica reglas de negocio específicas
para el formato de certificados (sangrías, interlineados).

Soporta:
    - Texto plano para variables normales
    - HTML enriquecido para campos específicos (OBJETIVO_PROGRAMA, CONTENIDO)

Author: Sistema de Certificados UNEMI
Version: 2.0.0
"""

import re
import logging
from typing import Dict, List, Optional, Set, Final
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches

# Configuración de logging
logger = logging.getLogger(__name__)

# ============================================================================
# CONSTANTES
# ============================================================================

# Variables que soportan contenido HTML enriquecido
# IMPORTANTE: Estas deben coincidir EXACTAMENTE con las variables en template_service.py
HTML_ENABLED_VARIABLES: Final[Set[str]] = {
    'OBJETIVO DEL PROGRAMA',  # ← Variable exacta usada en template_service
    'CONTENIDO',              # ← Variable exacta usada en template_service
    # Variantes alternativas por compatibilidad
    'OBJETIVO_PROGRAMA',
    'OBJETIVO',
    'CONTENIDO_PROGRAMA',
    'CONTENIDO DEL PROGRAMA'
}

# Configuración de formato para certificados
INDENT_SIZE_INCHES: Final[float] = 1.0  # Sangrías de 1.0 pulgadas
NAME_FONT_SIZE_PT: Final[int] = 22      # Tamaño de fuente para nombres
NAME_SPACE_BEFORE_PT: Final[int] = 24   # Espacio antes del nombre


# ============================================================================
# PROCESADOR DE FORMATO PARA CERTIFICADOS
# ============================================================================

class CertificadoPostProcessor:
    """
    Aplica reglas de formato específicas para los certificados de UNEMI.
    
    Responsabilidades:
        - Aplicar sangrías al cuerpo principal del certificado (primera página)
        - Aplicar formato especial al nombre del estudiante
        - Garantizar que las tablas no tengan sangrías
    
    Principios:
        - Patrón: Strategy (estrategia de formato específica para certificados)
        - SRP: Única responsabilidad de formato post-procesamiento
    """

    @staticmethod
    def apply_rules(doc: Document, variables: Dict[str, str]) -> None:
        """
        Aplica reglas de negocio al documento después del reemplazo.

        Args:
            doc: Documento de python-docx procesado.
            variables: Diccionario de variables usadas en el reemplazo.
            
        Returns:
            None (modifica el documento in-place)
        """
        # Extraer valores relevantes
        objetivo_value = variables.get('OBJETIVO DEL PROGRAMA', '') or variables.get('OBJETIVO_PROGRAMA', '')
        contenido_value = variables.get('CONTENIDO', '') or variables.get('CONTENIDO_PROGRAMA', '')
        nombres_value = variables.get('NOMBRES', '')

        # Paso 1: Analizar estructura del documento
        structure_map = CertificadoPostProcessor._analyze_document_structure(
            doc.paragraphs, objetivo_value, contenido_value
        )

        # Paso 2: Aplicar formato según la estructura
        CertificadoPostProcessor._apply_formatting(
            doc.paragraphs, structure_map, nombres_value
        )
        
        # Paso 3: Limpiar formato heredado en tablas
        CertificadoPostProcessor._clean_table_formatting(doc.tables)

    @staticmethod
    def _apply_formatting(
        paragraphs: List[Paragraph],
        structure_map: Dict[int, str],
        nombres_value: str
    ) -> None:
        """
        Aplica formato a los párrafos según su tipo en el mapa de estructura.
        
        Args:
            paragraphs: Lista de párrafos del documento.
            structure_map: Mapa de índice -> tipo de párrafo.
            nombres_value: Valor del nombre del estudiante.
        """
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            paragraph_type = structure_map.get(i, 'skip')
            
            # Aplicar sangrías a párrafos del cuerpo
            if paragraph_type == 'body':
                try:
                    paragraph.paragraph_format.left_indent = Inches(INDENT_SIZE_INCHES)
                    paragraph.paragraph_format.right_indent = Inches(INDENT_SIZE_INCHES)
                except Exception as e:
                    logger.warning(f"No se pudo aplicar sangría en párrafo '{text[:20]}...': {e}")

            # Formato especial para el nombre del estudiante
            if nombres_value and nombres_value in text and len(text) < 100:
                try:
                    paragraph.paragraph_format.space_before = Pt(NAME_SPACE_BEFORE_PT)
                    for run in paragraph.runs:
                        if run.font:
                            run.font.size = Pt(NAME_FONT_SIZE_PT)
                except Exception as e:
                    logger.warning(f"No se pudo aplicar formato al nombre: {e}")

    @staticmethod
    def _clean_table_formatting(tables: List) -> None:
        """
        Remueve sangrías heredadas de las tablas.
        
        Args:
            tables: Lista de tablas del documento.
        """
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.paragraph_format.left_indent = None
                            paragraph.paragraph_format.right_indent = None
                        except Exception:
                            pass  # Silenciar errores de formato

    @staticmethod
    def _analyze_document_structure(
        paragraphs: List[Paragraph],
        objetivo_value: str,
        contenido_value: str
    ) -> Dict[int, str]:
        """
        Analiza la estructura completa del documento y clasifica cada párrafo.
        
        ESTRATEGIA: Aplicar sangrías a TODO el texto que viene DESPUÉS 
        del nombre del estudiante, SOLO EN LA PRIMERA PÁGINA.
        
        Args:
            paragraphs: Lista de párrafos a analizar.
            objetivo_value: Contenido de OBJETIVO DEL PROGRAMA.
            contenido_value: Contenido de CONTENIDO.
        
        Returns:
            Dict con índice de párrafo -> tipo ('body', 'title', 'signature', 'table_content', 'skip')
        """
        structure_map: Dict[int, str] = {}
        found_student_name = False
        name_paragraph_index = -1
        end_of_first_page = False
        
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                structure_map[i] = 'skip'
                continue
            
            # === DETECTAR EL NOMBRE DEL ESTUDIANTE ===
            if not found_student_name:
                if CertificadoPostProcessor._is_student_name(text):
                    found_student_name = True
                    name_paragraph_index = i
                    structure_map[i] = 'skip'
                    continue
            
            # === EXCLUSIONES CRÍTICAS ===
            
            # 1. Contenido de tablas (OBJETIVO/CONTENIDO)
            if CertificadoPostProcessor._is_table_content(text, objetivo_value, contenido_value):
                structure_map[i] = 'table_content'
                end_of_first_page = True
                continue
            
            # 2. Títulos de sección
            if CertificadoPostProcessor._is_section_title(text):
                structure_map[i] = 'title'
                end_of_first_page = True
                continue
            
            # 3. Líneas de firma
            if CertificadoPostProcessor._is_signature(text):
                structure_map[i] = 'signature'
                continue
            
            # 4. Títulos de documento
            if CertificadoPostProcessor._is_document_title(text):
                structure_map[i] = 'title'
                continue
            
            # === REGLA PRINCIPAL ===
            # Aplicar a todo después del nombre, solo en primera página
            if found_student_name and i > name_paragraph_index and not end_of_first_page:
                structure_map[i] = 'body'
            else:
                structure_map[i] = 'skip'
        
        return structure_map

    @staticmethod
    def _is_student_name(text: str) -> bool:
        """Detecta si un párrafo contiene el nombre del estudiante."""
        return (
            text.isupper() and 
            10 < len(text) < 100 and 
            'CERTIFICADO' not in text and 
            'OBJETIVO' not in text and 
            'CONTENIDO' not in text
        )

    @staticmethod
    def _is_table_content(text: str, objetivo_value: str, contenido_value: str) -> bool:
        """Detecta si un párrafo contiene contenido de tabla."""
        if objetivo_value and len(objetivo_value) > 20 and objetivo_value[:30] in text:
            return True
        if contenido_value and len(contenido_value) > 20 and contenido_value[:30] in text:
            return True
        return False

    @staticmethod
    def _is_section_title(text: str) -> bool:
        """Detecta si un párrafo es un título de sección."""
        section_titles = [
            'Objetivo del programa', 'OBJETIVO DEL PROGRAMA', 'OBJETIVO',
            'Contenido del programa', 'CONTENIDO DEL PROGRAMA', 'CONTENIDO',
            'Modalidad:', 'MODALIDAD:', 'Duración:', 'DURACIÓN:'
        ]
        return any(title in text for title in section_titles)

    @staticmethod
    def _is_signature(text: str) -> bool:
        """Detecta si un párrafo es una línea de firma."""
        signature_prefixes = ('Ph.', 'Msc.', 'Ing.', 'Lic.', 'Dr.', 'Dra.')
        signature_keywords = ['Rector', 'Director', 'Decano', 'Coordinador', 
                            'Secretario', 'Vicerrector', 'Vicedecano']
        return (
            text.startswith(signature_prefixes) or 
            any(kw in text for kw in signature_keywords)
        )

    @staticmethod
    def _is_document_title(text: str) -> bool:
        """Detecta si un párrafo es un título del documento."""
        return len(text) < 30 and text.isupper() and 'CERTIFICADO' in text


# ============================================================================
# MOTOR DE REEMPLAZO DE VARIABLES
# ============================================================================

class VariableReplacer:
    """
    Motor de reemplazo de variables en documentos Word.
    
    Características:
        - Preserva estrictamente el formato original (fuentes, negritas, colores)
        - Soporta texto plano para variables estándar
        - Soporta HTML enriquecido para OBJETIVO_PROGRAMA y CONTENIDO
        - Maneja fragmentación de runs en Word
    
    Patrón: Facade (interfaz simplificada para reemplazo complejo)
    """

    @classmethod
    def process(cls, doc_path: str, variables: Dict[str, str]) -> Document:
        """
        Ejecuta el proceso completo de reemplazo de variables.

        Args:
            doc_path: Ruta al archivo .docx plantilla.
            variables: Diccionario {clave: valor} a reemplazar.

        Returns:
            Documento procesado (listo para guardar).
            
        Raises:
            Exception: Si hay un error crítico procesando la plantilla.
        """
        try:
            doc = Document(doc_path)
            normalized_vars = cls._normalize_variables(variables)
            
            # Reemplazo en todas las secciones del documento
            cls._replace_in_paragraphs(doc.paragraphs, normalized_vars, doc)
            cls._replace_in_tables(doc.tables, normalized_vars, doc)
            cls._replace_in_headers_footers(doc.sections, normalized_vars, doc)
                
            # Aplicar reglas de formato de certificados
            CertificadoPostProcessor.apply_rules(doc, normalized_vars)
            
            return doc
            
        except Exception as e:
            logger.error(f"Error crítico procesando plantilla {doc_path}: {e}", exc_info=True)
            raise

    @staticmethod
    def _normalize_variables(variables: Dict[str, str]) -> Dict[str, str]:
        """
        Normaliza claves de variables para soportar variantes.
        
        Args:
            variables: Diccionario original de variables.
        
        Returns:
            Diccionario normalizado con variantes de espacios y guiones bajos.
        
        Ejemplo:
            'NOMBRE COMPLETO' -> también crea 'NOMBRE_COMPLETO'
        """
        result = {}
        for k, v in variables.items():
            key_upper = k.upper()
            result[key_upper] = str(v)
            
            # Crear variantes para facilitar coincidencias
            if ' ' in key_upper:
                result[key_upper.replace(' ', '_')] = str(v)
            elif '_' in key_upper:
                result[key_upper.replace('_', ' ')] = str(v)
        
        return result

    @classmethod
    def _replace_in_tables(cls, tables: List, normalized_vars: Dict[str, str], document: Document) -> None:
        """Reemplaza variables en todas las tablas del documento."""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cls._replace_in_paragraphs(cell.paragraphs, normalized_vars, document)

    @classmethod
    def _replace_in_headers_footers(
        cls, 
        sections: List, 
        normalized_vars: Dict[str, str], 
        document: Document
    ) -> None:
        """Reemplaza variables en encabezados y pies de página."""
        for section in sections:
            cls._replace_in_paragraphs(section.header.paragraphs, normalized_vars, document)
            cls._replace_in_paragraphs(section.footer.paragraphs, normalized_vars, document)

    @classmethod
    def _replace_in_paragraphs(
        cls, 
        paragraphs: List[Paragraph], 
        variables: Dict[str, str], 
        document: Optional[Document] = None
    ) -> None:
        """
        Itera sobre párrafos y delega el reemplazo.
        
        Args:
            paragraphs: Lista de párrafos a procesar.
            variables: Variables a reemplazar.
            document: Documento padre (necesario para HTML).
        """
        for paragraph in paragraphs:
            cls._replace_in_single_paragraph(paragraph, variables, document)

    @classmethod
    def _replace_in_single_paragraph(
        cls, 
        paragraph: Paragraph, 
        variables: Dict[str, str], 
        document: Optional[Document] = None
    ) -> None:
        """
        Busca y reemplaza {{VAR}} en un párrafo, manejando runs fragmentados.
        
        Args:
            paragraph: Párrafo a procesar.
            variables: Variables a reemplazar.
            document: Documento padre (para conversión HTML).
        """
        # Optimización: chequeo rápido
        full_text = paragraph.text
        if '{' not in full_text:
            return

        # Regex para encontrar {{CUALQUIER_COSA}}
        pattern = r'\{\{[A-ZÁÉÍÓÚÑa-záéíóúñ0-9_ ]+\}\}'
        
        while True:
            current_text = paragraph.text
            match = re.search(pattern, current_text)
            if not match:
                break
                
            var_placeholder = match.group()
            var_name = var_placeholder[2:-2].strip().upper()
            
            # Resolver valor
            value = cls._resolve_value(var_name, variables)
            if value is None:
                break
            
            # Detectar si es una variable HTML
            if cls._is_html_variable(var_name) and cls._contains_html_tags(value):
                cls._replace_with_html(paragraph, match.start(), match.end(), value, document)
            else:
                cls._replace_run_content(paragraph, match.start(), match.end(), value)
    
    @staticmethod
    def _is_html_variable(var_name: str) -> bool:
        """Verifica si una variable soporta HTML enriquecido."""
        return var_name.upper() in HTML_ENABLED_VARIABLES
    
    @staticmethod
    def _contains_html_tags(content: str) -> bool:
        """Detecta si el contenido tiene tags HTML."""
        if not content:
            return False
        html_pattern = r'<(p|br|strong|b|em|i|u|ul|ol|li|table|tr|td|th|div|span)[\s>]'
        return bool(re.search(html_pattern, content, re.IGNORECASE))
    
    @staticmethod
    def _replace_with_html(
        paragraph: Paragraph, 
        start_idx: int, 
        end_idx: int, 
        html_content: str, 
        document: Optional[Document]
    ) -> None:
        """
        Reemplaza el marcador con contenido HTML convertido a Word.
        
        Args:
            paragraph: Párrafo que contiene la variable.
            start_idx: Índice de inicio del marcador.
            end_idx: Índice de fin del marcador.
            html_content: Contenido HTML a insertar.
            document: Documento padre.
        """
        try:
            from ..services.html_to_word_converter import HTMLToWordConverter
            
            # Limpiar el marcador del párrafo actual
            current_idx = 0
            for run in paragraph.runs:
                run_len = len(run.text)
                run_end = current_idx + run_len
                
                intersection_start = max(current_idx, start_idx)
                intersection_end = min(run_end, end_idx)
                
                if intersection_start < intersection_end:
                    local_start = intersection_start - current_idx
                    local_end = intersection_end - current_idx
                    
                    head = run.text[:local_start]
                    tail = run.text[local_end:]
                    run.text = head + tail
                
                current_idx += run_len
            
            # Convertir HTML a Word e insertar
            converter = HTMLToWordConverter()
            converter.convert_and_insert(html_content, paragraph, clear_paragraph=False, document=document)
            
        except Exception as e:
            logger.error(f"Error al convertir HTML a Word: {e}", exc_info=True)
            # Fallback: insertar como texto plano
            VariableReplacer._replace_run_content(paragraph, start_idx, end_idx, html_content)

    @staticmethod
    def _resolve_value(var_name: str, variables: Dict[str, str]) -> Optional[str]:
        """
        Busca el valor de la variable intentando varias claves.
        
        Args:
            var_name: Nombre de la variable a resolver.
            variables: Diccionario de variables.
        
        Returns:
            Valor de la variable o None si no se encuentra.
        """
        # Intento directo
        val = variables.get(var_name)
        if val is not None:
            # Transformación específica para NOMBRES
            if 'NOMBRE' in var_name and 'CURSO' not in var_name and 'EVENTO' not in var_name:
                return val.upper()
            return val
            
        # Intentar variantes
        val = variables.get(var_name.replace(' ', '_'))
        if val is not None:
            return val
        
        val = variables.get(var_name.replace('_', ' '))
        if val is not None:
            return val
        
        return None

    @staticmethod
    def _replace_run_content(
        paragraph: Paragraph, 
        start_idx: int, 
        end_idx: int, 
        replacement: str
    ) -> None:
        """
        Reemplaza texto en el rango [start_idx, end_idx) preservando formato.
        
        Args:
            paragraph: Párrafo que contiene el texto.
            start_idx: Índice de inicio del reemplazo.
            end_idx: Índice de fin del reemplazo.
            replacement: Texto nuevo.
        """
        current_idx = 0
        first_run_index = -1
        runs = paragraph.runs
        
        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_end = current_idx + run_len
            
            intersection_start = max(current_idx, start_idx)
            intersection_end = min(run_end, end_idx)
            
            if intersection_start < intersection_end:
                local_start = intersection_start - current_idx
                local_end = intersection_end - current_idx
                
                head = run.text[:local_start]
                tail = run.text[local_end:]
                
                if first_run_index == -1:
                    first_run_index = i
                    run.text = head + replacement + tail
                else:
                    run.text = head + tail
            
            current_idx += run_len


# ============================================================================
# FUNCIÓN PÚBLICA DE COMPATIBILIDAD
# ============================================================================

def replace_variables_in_template(template_path: str, variables: Dict[str, str]) -> Document:
    """
    Función helper pública para mantener compatibilidad con código existente.
    
    Args:
        template_path: Ruta al archivo .docx plantilla.
        variables: Diccionario de variables a reemplazar.
    
    Returns:
        Documento procesado.
    """
    return VariableReplacer.process(template_path, variables)
