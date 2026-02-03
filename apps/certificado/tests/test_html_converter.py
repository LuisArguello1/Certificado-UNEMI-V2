"""
Tests unitarios para el conversor HTML a Word.

Verifica que la conversión de HTML a formato Word nativo funcione correctamente
para todos los casos soportados: negritas, cursivas, listas, tablas, etc.

IMPORTANTE: El conversor mantiene TODO el contenido en el párrafo original
usando runs y saltos de línea. No crea nuevos párrafos en el documento.
"""

import unittest
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from docx.text.paragraph import Paragraph
from apps.certificado.services.html_to_word_converter import (
    HTMLToWordConverter,
    convert_html_to_word_inline,
    strip_html_tags
)


class TestHTMLToWordConverter(unittest.TestCase):
    """Tests para HTMLToWordConverter básico."""
    
    def setUp(self):
        """Configuración inicial para cada test."""
        self.converter = HTMLToWordConverter()
        self.doc = Document()
    
    def test_basic_bold_text(self):
        """Test: Conversión de texto en negrita."""
        html = "<p><strong>Texto en negrita</strong></p>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Verificar que el párrafo tiene contenido
        self.assertGreater(len(paragraph.runs), 0)
        
        # Verificar que hay al menos un run con negrita
        has_bold = any(run.bold for run in paragraph.runs if run.text.strip())
        self.assertTrue(has_bold, "Debe haber al menos un run en negrita")
    
    def test_basic_italic_text(self):
        """Test: Conversión de texto en cursiva."""
        html = "<p><em>Texto en cursiva</em></p>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        has_italic = any(run.italic for run in paragraph.runs if run.text.strip())
        self.assertTrue(has_italic, "Debe haber al menos un run en cursiva")
    
    def test_underline_text(self):
        """Test: Conversión de texto subrayado."""
        html = "<u>Texto subrayado</u>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        has_underline = any(run.underline for run in paragraph.runs if run.text.strip())
        self.assertTrue(has_underline, "Debe haber al menos un run subrayado")
    
    def test_combined_formatting(self):
        """Test: Combinación de formatos (negrita + cursiva)."""
        html = "<strong><em>Texto en negrita y cursiva</em></strong>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Verificar que hay runs con ambos formatos
        bold_italic_runs = [
            run for run in paragraph.runs 
            if run.bold and run.italic and run.text.strip()
        ]
        self.assertGreater(len(bold_italic_runs), 0, 
                          "Debe haber runs con negrita y cursiva")
    
    def test_empty_html(self):
        """Test: HTML vacío no debe causar error."""
        html = ""
        paragraph = self.doc.add_paragraph()
        
        # No debe lanzar excepción
        try:
            self.converter.convert_and_insert(html, paragraph)
        except Exception as e:
            self.fail(f"No debería lanzar excepción con HTML vacío: {e}")
    
    def test_plain_text_without_tags(self):
        """Test: Texto plano sin tags HTML."""
        html = "Texto simple sin formato"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Debe tener contenido
        self.assertIn("Texto simple", paragraph.text)
    
    def test_plain_text_with_line_breaks(self):
        """Test: Texto plano con saltos de línea \\r\\n."""
        text = "Primera línea\r\nSegunda línea\r\nTercera línea"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(text, paragraph)
        
        # Debe tener todo el contenido
        self.assertIn("Primera línea", paragraph.text)
        self.assertIn("Segunda línea", paragraph.text)
        self.assertIn("Tercera línea", paragraph.text)
    
    def test_plain_text_with_unix_line_breaks(self):
        """Test: Texto plano con saltos de línea \\n (Unix)."""
        text = "Línea uno\nLínea dos\nLínea tres"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(text, paragraph)
        
        self.assertIn("Línea uno", paragraph.text)
        self.assertIn("Línea dos", paragraph.text)
        self.assertIn("Línea tres", paragraph.text)
    
    def test_plain_text_multiline_real_content(self):
        """Test: Contenido real de texto plano multi-línea."""
        text = """1. Las ferias como herramienta efectiva de marketing
Importancia de la participación en ferias y eventos.
Diseño del stand y presentación del producto.
Estrategias de comunicación y promoción en ferias.
2. El seguimiento post-feria como estrategia de ventas
Técnicas para mantener el contacto con clientes potenciales."""
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(text, paragraph)
        
        # Todo el contenido debe estar presente
        self.assertIn("Las ferias como herramienta", paragraph.text)
        self.assertIn("Importancia de la participación", paragraph.text)
        self.assertIn("seguimiento post-feria", paragraph.text)

    def test_line_break(self):
        """Test: Saltos de línea <br>."""
        html = "Línea 1<br>Línea 2"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Verificar que se procesó
        self.assertGreater(len(paragraph.runs), 0)
        self.assertIn("Línea 1", paragraph.text)
        self.assertIn("Línea 2", paragraph.text)
    
    def test_unordered_list_inline(self):
        """Test: Lista desordenada <ul> - todo en un párrafo."""
        html = """
        <ul>
            <li>Item 1</li>
            <li>Item 2</li>
            <li>Item 3</li>
        </ul>
        """
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Todo debe estar en el mismo párrafo con viñetas
        self.assertIn("•", paragraph.text)
        self.assertIn("Item 1", paragraph.text)
        self.assertIn("Item 2", paragraph.text)
        self.assertIn("Item 3", paragraph.text)
    
    def test_ordered_list_inline(self):
        """Test: Lista ordenada <ol> - todo en un párrafo."""
        html = """
        <ol>
            <li>Primer punto</li>
            <li>Segundo punto</li>
        </ol>
        """
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Todo debe estar en el mismo párrafo con números
        self.assertIn("1.", paragraph.text)
        self.assertIn("2.", paragraph.text)
        self.assertIn("Primer punto", paragraph.text)
        self.assertIn("Segundo punto", paragraph.text)
    
    def test_table_inline(self):
        """Test: Tabla HTML - todo en un párrafo con tabulaciones."""
        html = """
        <table>
            <tr>
                <th>Encabezado 1</th>
                <th>Encabezado 2</th>
            </tr>
            <tr>
                <td>Celda 1</td>
                <td>Celda 2</td>
            </tr>
        </table>
        """
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Todo debe estar en el mismo párrafo
        self.assertIn("Encabezado 1", paragraph.text)
        self.assertIn("Celda 1", paragraph.text)
        self.assertIn("Celda 2", paragraph.text)
    
    def test_malformed_html(self):
        """Test: HTML mal formado no debe romper el conversor."""
        html = "<strong>Negrita sin cerrar"
        paragraph = self.doc.add_paragraph()
        
        # BeautifulSoup maneja HTML mal formado
        try:
            self.converter.convert_and_insert(html, paragraph)
        except Exception as e:
            self.fail(f"No debería fallar con HTML mal formado: {e}")
    
    def test_nested_tags(self):
        """Test: Tags anidados complejos."""
        html = "<p><strong>Negrita <em>y cursiva</em> dentro</strong></p>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Debe procesar correctamente
        self.assertGreater(len(paragraph.runs), 0)
        self.assertIn("Negrita", paragraph.text)
        self.assertIn("cursiva", paragraph.text)
    
    def test_multiple_paragraphs_stay_inline(self):
        """Test: Múltiples párrafos <p> se quedan en el mismo párrafo con saltos."""
        html = """
        <p>Primer párrafo</p>
        <p>Segundo párrafo</p>
        <p>Tercer párrafo</p>
        """
        paragraph = self.doc.add_paragraph()
        initial_paragraph_count = len(self.doc.paragraphs)
        
        self.converter.convert_and_insert(html, paragraph)
        
        # NO debe crear nuevos párrafos en el documento
        self.assertEqual(len(self.doc.paragraphs), initial_paragraph_count,
                        "No debe crear párrafos nuevos en el documento")
        
        # Todo el contenido debe estar en el mismo párrafo
        self.assertIn("Primer párrafo", paragraph.text)
        self.assertIn("Segundo párrafo", paragraph.text)
        self.assertIn("Tercer párrafo", paragraph.text)


class TestHelperFunctions(unittest.TestCase):
    """Tests para las funciones helper públicas."""
    
    def test_convert_html_to_word_inline_basic(self):
        """Test: Función helper básica."""
        doc = Document()
        paragraph = doc.add_paragraph()
        html = "<strong>Texto de prueba</strong>"
        
        # No debe lanzar error
        try:
            convert_html_to_word_inline(html, paragraph)
        except Exception as e:
            self.fail(f"Error en helper: {e}")
        
        # Verificar que procesó
        self.assertGreater(len(paragraph.runs), 0)
    
    def test_strip_html_tags(self):
        """Test: Función para eliminar tags HTML."""
        html = "<p><strong>Hola</strong> <em>mundo</em></p>"
        result = strip_html_tags(html)
        
        self.assertEqual(result.strip(), "Hola mundo")
    
    def test_strip_html_tags_empty(self):
        """Test: strip_html_tags con cadena vacía."""
        result = strip_html_tags("")
        self.assertEqual(result, "")
    
    def test_strip_html_tags_none(self):
        """Test: strip_html_tags con None."""
        result = strip_html_tags(None)
        self.assertEqual(result, "")


class TestEdgeCases(unittest.TestCase):
    """Tests para casos extremos."""
    
    def setUp(self):
        """Configuración inicial."""
        self.converter = HTMLToWordConverter()
        self.doc = Document()
    
    def test_very_long_text(self):
        """Test: Texto muy largo."""
        long_text = "Lorem ipsum " * 1000
        html = f"<p>{long_text}</p>"
        paragraph = self.doc.add_paragraph()
        
        try:
            self.converter.convert_and_insert(html, paragraph)
        except Exception as e:
            self.fail(f"Error con texto largo: {e}")
        
        # El texto debe estar en el párrafo
        self.assertIn("Lorem ipsum", paragraph.text)
    
    def test_special_characters(self):
        """Test: Caracteres especiales."""
        html = "<p>Áéíóú ñÑ ¿? ¡! @#$%</p>"
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Debe contener los caracteres especiales
        self.assertIn("Á", paragraph.text)
        self.assertIn("ñ", paragraph.text)
    
    def test_empty_table(self):
        """Test: Tabla vacía."""
        html = "<table></table>"
        paragraph = self.doc.add_paragraph()
        
        # No debe fallar
        try:
            self.converter.convert_and_insert(html, paragraph)
        except Exception as e:
            self.fail(f"Error con tabla vacía: {e}")
    
    def test_table_without_headers(self):
        """Test: Tabla sin encabezados."""
        html = """
        <table>
            <tr>
                <td>Dato 1</td>
                <td>Dato 2</td>
            </tr>
        </table>
        """
        paragraph = self.doc.add_paragraph()
        
        try:
            self.converter.convert_and_insert(html, paragraph)
            self.assertIn("Dato 1", paragraph.text)
            self.assertIn("Dato 2", paragraph.text)
        except Exception as e:
            self.fail(f"Error con tabla sin headers: {e}")
    
    def test_nested_lists(self):
        """Test: Listas anidadas."""
        html = """
        <ul>
            <li>Item 1
                <ul>
                    <li>Sub item 1</li>
                </ul>
            </li>
            <li>Item 2</li>
        </ul>
        """
        paragraph = self.doc.add_paragraph()
        
        try:
            self.converter.convert_and_insert(html, paragraph)
            self.assertIn("Item 1", paragraph.text)
            self.assertIn("Sub item", paragraph.text)
        except Exception as e:
            self.fail(f"Error con listas anidadas: {e}")


class TestRealWorldScenarios(unittest.TestCase):
    """Tests con escenarios reales del sistema de certificados."""
    
    def setUp(self):
        """Configuración inicial."""
        self.converter = HTMLToWordConverter()
        self.doc = Document()
    
    def test_objetivo_programa_real(self):
        """Test: Objetivo del programa con formato real."""
        html = """
        <p><strong>Objetivo General:</strong></p>
        <p>Capacitar a los participantes en:</p>
        <ul>
            <li>Gestión administrativa</li>
            <li>Liderazgo efectivo</li>
            <li>Comunicación asertiva</li>
        </ul>
        """
        paragraph = self.doc.add_paragraph()
        initial_count = len(self.doc.paragraphs)
        
        self.converter.convert_and_insert(html, paragraph)
        
        # NO debe crear párrafos nuevos
        self.assertEqual(len(self.doc.paragraphs), initial_count)
        
        # Todo el contenido debe estar en el mismo párrafo
        self.assertIn("Objetivo General:", paragraph.text)
        self.assertIn("Capacitar", paragraph.text)
        self.assertIn("Gestión administrativa", paragraph.text)
        self.assertIn("•", paragraph.text)  # Viñetas
    
    def test_contenido_programa_con_tabla(self):
        """Test: Contenido del programa con tabla de módulos."""
        html = """
        <p><strong>Contenido del Programa:</strong></p>
        <table>
            <tr>
                <th>Módulo</th>
                <th>Horas</th>
            </tr>
            <tr>
                <td>Introducción</td>
                <td>10h</td>
            </tr>
            <tr>
                <td>Práctica</td>
                <td>30h</td>
            </tr>
        </table>
        """
        paragraph = self.doc.add_paragraph()
        initial_count = len(self.doc.paragraphs)
        
        self.converter.convert_and_insert(html, paragraph)
        
        # NO debe crear párrafos nuevos
        self.assertEqual(len(self.doc.paragraphs), initial_count)
        
        # Todo debe estar en el mismo párrafo
        self.assertIn("Contenido del Programa:", paragraph.text)
        self.assertIn("Módulo", paragraph.text)
        self.assertIn("Horas", paragraph.text)
        self.assertIn("Introducción", paragraph.text)
        self.assertIn("10h", paragraph.text)
    
    def test_mixed_content_stays_inline(self):
        """Test: Contenido mixto (texto, lista, tabla) en un solo párrafo."""
        html = """
        <p><strong>Hola buenos días</strong></p>
        <p><em>como están todos</em></p>
        <ol>
            <li><em>Lista 1</em></li>
            <li><em>Lista 2</em></li>
            <li><em>Lista 3</em></li>
        </ol>
        <p>Probando</p>
        <table>
            <tr>
                <td>Probando</td>
                <td>probando</td>
            </tr>
        </table>
        """
        paragraph = self.doc.add_paragraph()
        initial_count = len(self.doc.paragraphs)
        
        self.converter.convert_and_insert(html, paragraph)
        
        # NO debe crear párrafos nuevos en el documento
        self.assertEqual(len(self.doc.paragraphs), initial_count,
                        "No debe crear párrafos nuevos - todo inline")
        
        # Todo el contenido debe estar en el mismo párrafo
        self.assertIn("Hola buenos días", paragraph.text)
        self.assertIn("como están todos", paragraph.text)
        self.assertIn("Lista 1", paragraph.text)
        self.assertIn("Lista 2", paragraph.text)
        self.assertIn("Lista 3", paragraph.text)
        self.assertIn("1.", paragraph.text)  # Numeración de lista
        self.assertIn("Probando", paragraph.text)


class TestFormattedListItems(unittest.TestCase):
    """Tests para items de lista con formato."""
    
    def setUp(self):
        """Configuración inicial."""
        self.converter = HTMLToWordConverter()
        self.doc = Document()
    
    def test_italic_list_items(self):
        """Test: Items de lista en cursiva."""
        html = """
        <ol>
            <li><em>Primer item</em></li>
            <li><em>Segundo item</em></li>
        </ol>
        """
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Debe haber runs con cursiva
        italic_runs = [r for r in paragraph.runs if r.italic and r.text.strip()]
        self.assertGreater(len(italic_runs), 0, "Debe haber items en cursiva")
    
    def test_bold_in_table(self):
        """Test: Celdas de tabla con negrita (th)."""
        html = """
        <table>
            <tr>
                <th>Header</th>
            </tr>
            <tr>
                <td>Normal</td>
            </tr>
        </table>
        """
        paragraph = self.doc.add_paragraph()
        
        self.converter.convert_and_insert(html, paragraph)
        
        # Debe haber runs con negrita (los th)
        bold_runs = [r for r in paragraph.runs if r.bold and r.text.strip()]
        self.assertGreater(len(bold_runs), 0, "Headers de tabla deben estar en negrita")


# Ejecutar tests
if __name__ == '__main__':
    unittest.main()
