"""
Servicio de conversión HTML a Word con formato nativo.

Este módulo convierte HTML generado por CKEditor a formato Word (.docx) nativo,
preservando el formato sin perder la estructura de la plantilla.

Soporta:
    - Negritas (<strong>, <b>)
    - Cursivas (<em>, <i>)
    - Subrayado (<u>)
    - Listas ordenadas y desordenadas (<ul>, <ol>, <li>) - como texto con viñetas
    - Saltos de línea (<br>, <p>) - como line breaks
    - Tablas (<table>, <tr>, <td>, <th>) - como tablas nativas de Word

Uso:
    >>> converter = HTMLToWordConverter()
    >>> doc = Document()
    >>> paragraph = doc.add_paragraph()
    >>> converter.convert_and_insert(html_content, paragraph, document=doc)
"""

from typing import Optional, List, Dict, Any, Tuple
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table
import logging
import re

logger = logging.getLogger(__name__)


class HTMLToWordConverter:
    """
    Conversor de HTML a formato Word nativo.
    
    Soporta tablas nativas de Word cuando se proporciona el documento.
    
    Attributes:
        default_font_size (int): Tamaño de fuente por defecto en puntos.
        preserve_spacing (bool): Si debe preservar espacios múltiples.
    """
    
    def __init__(
        self,
        default_font_size: int = 11,
        preserve_spacing: bool = True
    ):
        self.default_font_size = default_font_size
        self.preserve_spacing = preserve_spacing
        self._needs_line_break = False
        self._document = None  # Referencia al documento para crear tablas
        self._current_paragraph = None  # Párrafo actual
        self._pending_tables = []  # Tablas pendientes de insertar
        
    def convert_and_insert(
        self,
        html_content: str,
        target_paragraph,
        clear_paragraph: bool = True,
        document=None,
        align_left: bool = True
    ) -> None:
        """
        Convierte HTML y lo inserta en un párrafo de Word.
        
        Args:
            html_content: Contenido HTML o texto plano a convertir.
            target_paragraph: Párrafo de Word donde insertar el contenido.
            clear_paragraph: Si debe limpiar el párrafo antes de insertar.
            document: Documento de Word (necesario para crear tablas nativas).
            align_left: Si es True, alinea el párrafo a la izquierda para evitar
                       justificación forzada en textos cortos.
        """
        if not html_content or not html_content.strip():
            logger.warning("Contenido vacío recibido en convert_and_insert")
            return
        
        try:
            # Guardar referencia al documento
            self._document = document or self._get_document_from_paragraph(target_paragraph)
            self._current_paragraph = target_paragraph
            self._pending_tables = []
            
            # Limpiar el párrafo si es necesario
            if clear_paragraph:
                self._clear_paragraph(target_paragraph)
            
            # Alinear a la izquierda para evitar justificación forzada
            if align_left:
                try:
                    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass  # Ignorar si no se puede cambiar la alineación
            
            # Detectar si es texto plano o HTML
            if self._is_plain_text(html_content):
                self._insert_plain_text(target_paragraph, html_content)
            else:
                self._needs_line_break = False
                soup = BeautifulSoup(html_content, 'html.parser')
                self._process_elements(
                    soup,
                    target_paragraph,
                    {'bold': False, 'italic': False, 'underline': False}
                )
            
            # Insertar tablas pendientes después del párrafo
            self._insert_pending_tables()
            
        except Exception as e:
            logger.error(f"Error al convertir contenido a Word: {e}", exc_info=True)
            self._insert_plain_text(target_paragraph, html_content)
    
    def _get_document_from_paragraph(self, paragraph):
        """Obtiene el documento de Word desde un párrafo."""
        try:
            if hasattr(paragraph, 'part') and hasattr(paragraph.part, 'document'):
                return paragraph.part.document
            if hasattr(paragraph, '_parent'):
                parent = paragraph._parent
                if hasattr(parent, 'part') and hasattr(parent.part, 'document'):
                    return parent.part.document
            return None
        except Exception as e:
            logger.warning(f"Error obteniendo documento: {e}")
            return None
    
    def _insert_pending_tables(self):
        """Inserta las tablas pendientes después del párrafo actual."""
        if not self._pending_tables or not self._document:
            logger.debug(f"No hay tablas pendientes o documento: tables={len(self._pending_tables) if self._pending_tables else 0}, doc={self._document is not None}")
            return
        
        try:
            # Obtener el elemento XML del párrafo actual
            para_element = self._current_paragraph._element
            
            for table_data in self._pending_tables:
                # Crear la tabla de Word
                rows_data = table_data['rows']
                if not rows_data:
                    continue
                
                num_rows = len(rows_data)
                num_cols = max(len(row) for row in rows_data) if rows_data else 1
                
                logger.info(f"Creando tabla nativa de Word: {num_rows} filas x {num_cols} columnas")
                
                # Crear tabla
                table = self._document.add_table(rows=num_rows, cols=num_cols)
                
                # Intentar aplicar estilo, pero no fallar si no existe
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    # El estilo no existe en la plantilla, aplicar bordes manualmente
                    self._apply_table_borders(table)
                
                # Configurar AutoFit para que la tabla se ajuste al contenido
                table.autofit = True
                
                # Llenar la tabla
                for row_idx, row_data in enumerate(rows_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data.get('text', '')
                            
                            # Aplicar formato
                            if cell_data.get('bold') and cell.paragraphs:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                
                # Configurar ancho de tabla al 100% y columnas equitativas
                self._set_table_autofit(table, num_cols)
                
                # Mover la tabla después del párrafo actual
                table_element = table._tbl
                para_element.addnext(table_element)
                
                logger.debug(f"Tabla insertada exitosamente: {num_rows}x{num_cols}")
                
                # Actualizar para que la siguiente tabla vaya después de esta
                para_element = table_element
                
        except Exception as e:
            logger.error(f"Error insertando tablas: {e}", exc_info=True)
        
        self._pending_tables = []
    
    def _set_table_autofit(self, table, num_cols: int) -> None:
        """
        Configura la tabla para que ajuste automáticamente al contenido.
        Las columnas se dimensionan según el texto, evitando partir palabras.
        """
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Configurar tblW (ancho de tabla) a 100% (5000 = 100% en fifths of a percent)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:type'), 'pct')  # Porcentaje
            tblW.set(qn('w:w'), '5000')    # 100% (5000 = 100% en unidades de Word)
            
            # Configurar tblLayout a autofit para ajustar al contenido sin partir palabras
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'autofit')  # Autofit permite ajuste dinámico
            
            # Distribuir columnas equitativamente
            col_width_pct = 5000 // num_cols  # Dividir 100% entre columnas
            
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._tc.insert(0, tcPr)
                    
                    # Configurar ancho de celda
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:type'), 'pct')
                    tcW.set(qn('w:w'), str(col_width_pct))
                    
                    # Configurar noWrap para evitar partición de palabras
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is not None:
                        tcPr.remove(noWrap)
                    
                    # Configurar los párrafos de la celda
                    for para in cell.paragraphs:
                        try:
                            pPr = para._p.get_or_add_pPr()
                            
                            # Desactivar hyphenation (NO partir palabras con guión)
                            suppressAutoHyphens = pPr.find(qn('w:suppressAutoHyphens'))
                            if suppressAutoHyphens is None:
                                suppressAutoHyphens = OxmlElement('w:suppressAutoHyphens')
                                pPr.append(suppressAutoHyphens)
                            suppressAutoHyphens.set(qn('w:val'), '1')
                            
                        except Exception:
                            pass
            
        except Exception as e:
            logger.warning(f"No se pudo configurar ancho de tabla: {e}")
    
    def _apply_table_borders(self, table) -> None:
        """Aplica bordes a una tabla cuando el estilo 'Table Grid' no está disponible."""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            # Crear elemento de bordes
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Grosor del borde
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Negro
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
                
        except Exception as e:
            logger.warning(f"No se pudieron aplicar bordes a la tabla: {e}")
    
    def _is_plain_text(self, content: str) -> bool:
        """Detecta si el contenido es texto plano o HTML."""
        if not content:
            return True
        html_pattern = r'<(p|br|strong|b|em|i|u|ul|ol|li|table|tr|td|th|div|span|h[1-6]|figure)[>\s/]'
        return not bool(re.search(html_pattern, content, re.IGNORECASE))
    
    def _insert_plain_text(self, paragraph, text: str) -> None:
        """Inserta texto plano preservando saltos de línea."""
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            # Preservar líneas vacías agregando un salto de línea
            if line.strip():
                run = paragraph.add_run(line.strip())
                if self.default_font_size:
                    run.font.size = Pt(self.default_font_size)
            # Siempre agregar break entre líneas (incluso si la línea está vacía)
            if i < len(lines) - 1:
                paragraph.add_run().add_break()
    
    def _clear_paragraph(self, paragraph) -> None:
        """Limpia todo el contenido de un párrafo."""
        for run in paragraph.runs:
            run.text = ''
    
    def _process_elements(
        self,
        element,
        paragraph,
        parent_formatting: Dict[str, bool],
        is_first_in_block: bool = True
    ) -> None:
        """Procesa recursivamente elementos HTML."""
        first_child = True
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if not self.preserve_spacing:
                    text = ' '.join(text.split())
                
                if text.strip():
                    if self._needs_line_break:
                        paragraph.add_run().add_break()
                        self._needs_line_break = False
                    self._add_formatted_text(paragraph, text, parent_formatting)
            
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                new_formatting = parent_formatting.copy()
                
                # === FORMATOS DE TEXTO ===
                if tag_name in ['strong', 'b']:
                    new_formatting['bold'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                elif tag_name in ['em', 'i']:
                    new_formatting['italic'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                elif tag_name == 'u':
                    new_formatting['underline'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                # === SALTOS DE LÍNEA ===
                elif tag_name == 'br':
                    paragraph.add_run().add_break()
                    self._needs_line_break = False
                
                elif tag_name == 'p':
                    # Siempre agregar salto de línea antes de un nuevo párrafo (excepto el primero)
                    if not is_first_in_block and not first_child:
                        paragraph.add_run().add_break()
                    
                    # Detectar alineación desde el atributo style o class de CKEditor
                    alignment = self._get_alignment_from_element(child)
                    if alignment:
                        try:
                            paragraph.alignment = alignment
                        except Exception:
                            pass
                    
                    # Detectar y aplicar line-height (interlineado)
                    line_height = self._get_line_height_from_element(child)
                    if line_height:
                        self._apply_line_height(paragraph, line_height)
                    
                    # Verificar si el párrafo está vacío (solo tiene &nbsp; o espacios)
                    child_text = child.get_text()
                    is_empty_paragraph = not child_text.strip() or child_text.strip() == '\u00a0'
                    
                    if is_empty_paragraph:
                        # Párrafo vacío: agregar un salto de línea para preservar el espacio
                        paragraph.add_run().add_break()
                    else:
                        self._process_elements(child, paragraph, parent_formatting, True)
                    
                    self._needs_line_break = True
                
                # === LISTAS ===
                elif tag_name in ['ul', 'ol']:
                    ordered = (tag_name == 'ol')
                    self._process_list_inline(child, paragraph, parent_formatting, ordered)
                
                # === TABLAS (nativas de Word) ===
                elif tag_name == 'table':
                    self._queue_table(child)
                
                # === FIGURE (CKEditor envuelve tablas) ===
                elif tag_name == 'figure':
                    table = child.find('table')
                    if table:
                        self._queue_table(table)
                    else:
                        self._process_elements(child, paragraph, parent_formatting, first_child)
                
                # === TBODY, THEAD ===
                elif tag_name in ['tbody', 'thead', 'tfoot']:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
                
                # === OTROS ===
                elif tag_name in ['div', 'span', 'a']:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
                
                elif tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    if not first_child:
                        paragraph.add_run().add_break()
                    
                    # Detectar alineación
                    alignment = self._get_alignment_from_element(child)
                    if alignment:
                        try:
                            paragraph.alignment = alignment
                        except Exception:
                            pass
                    
                    new_formatting['bold'] = True
                    # Tamaño según nivel de heading
                    heading_sizes = {'h1': 18, 'h2': 16, 'h3': 14, 'h4': 12, 'h5': 11, 'h6': 11}
                    heading_size = heading_sizes.get(tag_name, self.default_font_size)
                    
                    # Procesar contenido con formato especial
                    for heading_child in child.children:
                        if isinstance(heading_child, NavigableString):
                            text = str(heading_child).strip()
                            if text:
                                run = paragraph.add_run(text)
                                run.bold = True
                                run.font.size = Pt(heading_size)
                        elif isinstance(heading_child, Tag):
                            # Procesar tags internos manteniendo el tamaño
                            text = heading_child.get_text(strip=True)
                            if text:
                                run = paragraph.add_run(text)
                                run.bold = True
                                run.font.size = Pt(heading_size)
                    
                    self._needs_line_break = True
                
                else:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
            
            first_child = False
    
    def _add_formatted_text(self, paragraph, text: str, formatting: Dict[str, bool]) -> None:
        """Añade texto con formato al párrafo."""
        run = paragraph.add_run(text)
        run.bold = formatting.get('bold', False)
        run.italic = formatting.get('italic', False)
        run.underline = formatting.get('underline', False)
        if self.default_font_size:
            run.font.size = Pt(self.default_font_size)
    
    def _get_alignment_from_element(self, element: Tag) -> Optional[Any]:
        """
        Extrae la alineación de un elemento HTML.
        
        Busca en:
        - Atributo class (text-left, text-center, text-right, text-justify)
        - Atributo style (text-align: left/center/right/justify)
        
        Returns:
            Constante WD_ALIGN_PARAGRAPH o None si no hay alineación.
        """
        try:
            # Buscar en las clases
            classes = element.get('class', [])
            if isinstance(classes, str):
                classes = classes.split()
            
            for cls in classes:
                if 'text-left' in cls or 'align-left' in cls:
                    return WD_ALIGN_PARAGRAPH.LEFT
                elif 'text-center' in cls or 'align-center' in cls:
                    return WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-right' in cls or 'align-right' in cls:
                    return WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-justify' in cls or 'align-justify' in cls:
                    return WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Buscar en el atributo style
            style = element.get('style', '')
            if 'text-align' in style:
                if 'left' in style:
                    return WD_ALIGN_PARAGRAPH.LEFT
                elif 'center' in style:
                    return WD_ALIGN_PARAGRAPH.CENTER
                elif 'right' in style:
                    return WD_ALIGN_PARAGRAPH.RIGHT
                elif 'justify' in style:
                    return WD_ALIGN_PARAGRAPH.JUSTIFY
        
        except Exception as e:
            logger.debug(f"Error detectando alineación: {e}")
        
        return None
    
    def _get_line_height_from_element(self, element: Tag) -> Optional[float]:
        """
        Extrae el interlineado de un elemento HTML.
        
        Busca en el atributo style (line-height: 1.5, etc.)
        
        Returns:
            Float con el valor de line-height o None.
        """
        try:
            style = element.get('style', '')
            if 'line-height' in style:
                # Extraer el valor numérico
                import re
                match = re.search(r'line-height:\s*([0-9.]+)', style)
                if match:
                    return float(match.group(1))
        except Exception as e:
            logger.debug(f"Error detectando line-height: {e}")
        
        return None
    
    def _apply_line_height(self, paragraph, line_height: float) -> None:
        """Aplica el interlineado a un párrafo de Word."""
        try:
            from docx.shared import Pt
            # Convertir line-height a espaciado de Word
            # line-height 1.0 = single, 1.5 = 1.5 lines, 2.0 = double
            if line_height <= 1.0:
                paragraph.paragraph_format.line_spacing = 1.0
            else:
                paragraph.paragraph_format.line_spacing = line_height
        except Exception as e:
            logger.debug(f"Error aplicando line-height: {e}")

    def _process_list_inline(
        self,
        list_element: Tag,
        paragraph,
        parent_formatting: Dict[str, bool],
        ordered: bool = False
    ) -> None:
        """Procesa listas HTML como texto con viñetas/números."""
        try:
            items = list_element.find_all('li', recursive=False)
            if not items:
                return
            
            # Añadir salto de línea antes de la lista si hay contenido previo
            if paragraph.text.strip():
                paragraph.add_run().add_break()
            
            for index, item in enumerate(items, start=1):
                # Añadir salto entre items (excepto el primero)
                if index > 1:
                    paragraph.add_run().add_break()
                
                # Prefijo de viñeta o número
                prefix = f"{index}. " if ordered else "• "
                prefix_run = paragraph.add_run(prefix)
                if self.default_font_size:
                    prefix_run.font.size = Pt(self.default_font_size)
                
                # Procesar contenido del item de forma especial para listas
                self._process_list_item_content(item, paragraph, parent_formatting)
            
            self._needs_line_break = True
            
        except Exception as e:
            logger.error(f"Error procesando lista: {e}", exc_info=True)
    
    def _process_list_item_content(
        self,
        item_element: Tag,
        paragraph,
        parent_formatting: Dict[str, bool]
    ) -> None:
        """
        Procesa el contenido de un elemento <li> sin añadir breaks entre elementos.
        Esto evita que los elementos internos (como <p>) rompan el formato de la lista.
        """
        for child in item_element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    self._add_formatted_text(paragraph, text, parent_formatting)
            
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                new_formatting = parent_formatting.copy()
                
                # Manejar formatos inline
                if tag_name in ['strong', 'b']:
                    new_formatting['bold'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name in ['em', 'i']:
                    new_formatting['italic'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name == 'u':
                    new_formatting['underline'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name == 'br':
                    paragraph.add_run().add_break()
                
                elif tag_name == 'p':
                    # Para <p> dentro de <li>, solo extraer el texto sin breaks adicionales
                    self._process_list_item_content(child, paragraph, parent_formatting)
                
                elif tag_name in ['span', 'a', 'div']:
                    self._process_list_item_content(child, paragraph, parent_formatting)
                
                elif tag_name in ['ul', 'ol']:
                    # Lista anidada - añadir con indentación
                    paragraph.add_run().add_break()
                    nested_ordered = (tag_name == 'ol')
                    nested_items = child.find_all('li', recursive=False)
                    for ni, nested_item in enumerate(nested_items, start=1):
                        if ni > 1:
                            paragraph.add_run().add_break()
                        nested_prefix = f"  {ni}. " if nested_ordered else "  ○ "
                        nested_run = paragraph.add_run(nested_prefix)
                        if self.default_font_size:
                            nested_run.font.size = Pt(self.default_font_size)
                        self._process_list_item_content(nested_item, paragraph, parent_formatting)
                
                else:
                    # Otros elementos: extraer texto
                    text = child.get_text(strip=True)
                    if text:
                        self._add_formatted_text(paragraph, text, parent_formatting)
    
    def _queue_table(self, table_element: Tag) -> None:
        """
        Encola una tabla para insertar después del párrafo.
        
        Extrae los datos de la tabla HTML y los guarda para crear
        una tabla nativa de Word después.
        """
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            table_data = {'rows': []}
            
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_data = []
                
                for cell in cells:
                    cell_text = cell.get_text(strip=True)
                    is_header = cell.name == 'th'
                    row_data.append({
                        'text': cell_text,
                        'bold': is_header
                    })
                
                if row_data:
                    table_data['rows'].append(row_data)
            
            if table_data['rows']:
                self._pending_tables.append(table_data)
                
        except Exception as e:
            logger.error(f"Error encolando tabla: {e}", exc_info=True)


# ============================================================================
# FUNCIONES AUXILIARES
# ============================================================================

def convert_html_to_word_inline(html_content: str, paragraph, document=None) -> None:
    """Función helper para convertir HTML a Word."""
    converter = HTMLToWordConverter()
    converter.convert_and_insert(html_content, paragraph, clear_paragraph=False, document=document)


def strip_html_tags(html_content: str) -> str:
    """Elimina tags HTML y devuelve solo el texto."""
    if not html_content:
        return ''
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text()
    except Exception:
        return html_content
