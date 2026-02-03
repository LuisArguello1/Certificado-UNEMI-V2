"""
Servicio de conversión HTML a Word con formato nativo.

Este módulo convierte HTML generado por CKEditor a formato Word (.docx) nativo,
preservando el formato sin perder la estructura de la plantilla.

Soporta:
    - Negritas (<strong>, <b>)
    - Cursivas (<em>, <i>)
    - Subrayado (<u>)
    - Listas ordenadas y desordenadas (<ul>, <ol>, <li>) - como texto con viñetas
    - Saltos de línea (<br>, <p>) - como line breaks
    - Tablas (<table>, <tr>, <td>, <th>) - como tablas nativas de Word

Uso:
    >>> converter = HTMLToWordConverter()
    >>> doc = Document()
    >>> paragraph = doc.add_paragraph()
    >>> converter.convert_and_insert(html_content, paragraph, document=doc)
"""

from typing import Optional, List, Dict, Any, Tuple
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table
import logging
import re

logger = logging.getLogger(__name__)


class HTMLToWordConverter:
    """
    Conversor de HTML a formato Word nativo.
    
    Soporta tablas nativas de Word cuando se proporciona el documento.
    
    Attributes:
        default_font_size (int): Tamaño de fuente por defecto en puntos.
        preserve_spacing (bool): Si debe preservar espacios múltiples.
    """
    
    def __init__(
        self,
        default_font_size: int = 11,
        preserve_spacing: bool = True
    ):
        self.default_font_size = default_font_size
        self.preserve_spacing = preserve_spacing
        self._needs_line_break = False
        self._document = None
        self._current_paragraph = None
        self._pending_tables = []
        
    def convert_and_insert(
        self,
        html_content: str,
        target_paragraph,
        clear_paragraph: bool = True,
        document=None,
        align_left: bool = True
    ) -> None:
        """
        Convierte HTML y lo inserta en un párrafo de Word.
        """
        if not html_content or not html_content.strip():
            logger.warning("Contenido vacío recibido en convert_and_insert")
            return
        
        try:
            self._document = document or self._get_document_from_paragraph(target_paragraph)
            self._current_paragraph = target_paragraph
            self._pending_tables = []
            
            if clear_paragraph:
                self._clear_paragraph(target_paragraph)
            
            if align_left:
                try:
                    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass
            
            if self._is_plain_text(html_content):
                self._insert_plain_text(target_paragraph, html_content)
            else:
                self._needs_line_break = False
                soup = BeautifulSoup(html_content, 'html.parser')
                self._process_elements(
                    soup,
                    target_paragraph,
                    {'bold': False, 'italic': False, 'underline': False}
                )
            
            self._insert_pending_tables()
            
        except Exception as e:
            logger.error(f"Error al convertir contenido a Word: {e}", exc_info=True)
            self._insert_plain_text(target_paragraph, html_content)
    
    def _get_document_from_paragraph(self, paragraph):
        """Obtiene el documento de Word desde un párrafo."""
        try:
            if hasattr(paragraph, 'part') and hasattr(paragraph.part, 'document'):
                return paragraph.part.document
            if hasattr(paragraph, '_parent'):
                parent = paragraph._parent
                if hasattr(parent, 'part') and hasattr(parent.part, 'document'):
                    return parent.part.document
            return None
        except Exception as e:
            logger.warning(f"Error obteniendo documento: {e}")
            return None
    
    def _insert_pending_tables(self):
        """Inserta las tablas pendientes después del párrafo actual."""
        if not self._pending_tables or not self._document:
            return
        
        try:
            para_element = self._current_paragraph._element
            
            for table_data in self._pending_tables:
                rows_data = table_data['rows']
                if not rows_data:
                    continue
                
                num_rows = len(rows_data)
                num_cols = max(len(row) for row in rows_data) if rows_data else 1
                
                logger.info(f"Creando tabla nativa de Word: {num_rows} filas x {num_cols} columnas")
                
                table = self._document.add_table(rows=num_rows, cols=num_cols)
                
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    self._apply_table_borders(table)
                
                table.autofit = True
                
                for row_idx, row_data in enumerate(rows_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data.get('text', '')
                            
                            if cell_data.get('bold') and cell.paragraphs:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                
                self._set_table_autofit(table, num_cols)
                
                table_element = table._tbl
                para_element.addnext(table_element)
                para_element = table_element
                
        except Exception as e:
            logger.error(f"Error insertando tablas: {e}", exc_info=True)
        
        self._pending_tables = []
    
    def _set_table_autofit(self, table, num_cols: int) -> None:
        """Configura la tabla para que use el 100% del ancho disponible."""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')
            
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed')
            
            col_width_pct = 5000 // num_cols
            
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._tc.insert(0, tcPr)
                    
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:type'), 'pct')
                    tcW.set(qn('w:w'), str(col_width_pct))
            
        except Exception as e:
            logger.warning(f"No se pudo configurar ancho de tabla: {e}")
    
    def _apply_table_borders(self, table) -> None:
        """Aplica bordes a una tabla cuando el estilo 'Table Grid' no está disponible."""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
                
        except Exception as e:
            logger.warning(f"No se pudieron aplicar bordes a la tabla: {e}")
    
    def _is_plain_text(self, content: str) -> bool:
        """Detecta si el contenido es texto plano o HTML."""
        if not content:
            return True
        html_pattern = r'<(p|br|strong|b|em|i|u|ul|ol|li|table|tr|td|th|div|span|h[1-6]|figure)[>\s/]'
        return not bool(re.search(html_pattern, content, re.IGNORECASE))
    
    def _insert_plain_text(self, paragraph, text: str) -> None:
        """Inserta texto plano preservando saltos de línea."""
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if line.strip():
                run = paragraph.add_run(line.strip())
                if self.default_font_size:
                    run.font.size = Pt(self.default_font_size)
            if i < len(lines) - 1:
                paragraph.add_run().add_break()
    
    def _clear_paragraph(self, paragraph) -> None:
        """Limpia todo el contenido de un párrafo."""
        for run in paragraph.runs:
            run.text = ''
    
    def _process_elements(
        self,
        element,
        paragraph,
        parent_formatting: Dict[str, bool],
        is_first_in_block: bool = True
    ) -> None:
        """Procesa recursivamente elementos HTML."""
        first_child = True
        
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if not self.preserve_spacing:
                    text = ' '.join(text.split())
                
                if text.strip():
                    if self._needs_line_break:
                        paragraph.add_run().add_break()
                        self._needs_line_break = False
                    self._add_formatted_text(paragraph, text, parent_formatting)
            
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                new_formatting = parent_formatting.copy()
                
                if tag_name in ['strong', 'b']:
                    new_formatting['bold'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                elif tag_name in ['em', 'i']:
                    new_formatting['italic'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                elif tag_name == 'u':
                    new_formatting['underline'] = True
                    self._process_elements(child, paragraph, new_formatting, first_child)
                
                elif tag_name == 'br':
                    paragraph.add_run().add_break()
                    self._needs_line_break = False
                
                elif tag_name == 'p':
                    if not is_first_in_block and not first_child:
                        paragraph.add_run().add_break()
                    self._process_elements(child, paragraph, parent_formatting, True)
                    self._needs_line_break = True
                
                elif tag_name in ['ul', 'ol']:
                    ordered = (tag_name == 'ol')
                    self._process_list_inline(child, paragraph, parent_formatting, ordered)
                
                elif tag_name == 'table':
                    self._queue_table(child)
                
                elif tag_name == 'figure':
                    table = child.find('table')
                    if table:
                        self._queue_table(table)
                    else:
                        self._process_elements(child, paragraph, parent_formatting, first_child)
                
                elif tag_name in ['tbody', 'thead', 'tfoot']:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
                
                elif tag_name in ['div', 'span', 'a']:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
                
                elif tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    if not first_child:
                        paragraph.add_run().add_break()
                    new_formatting['bold'] = True
                    self._process_elements(child, paragraph, new_formatting, True)
                    self._needs_line_break = True
                
                else:
                    self._process_elements(child, paragraph, parent_formatting, first_child)
            
            first_child = False
    
    def _add_formatted_text(self, paragraph, text: str, formatting: Dict[str, bool]) -> None:
        """Añade texto con formato al párrafo."""
        run = paragraph.add_run(text)
        run.bold = formatting.get('bold', False)
        run.italic = formatting.get('italic', False)
        run.underline = formatting.get('underline', False)
        if self.default_font_size:
            run.font.size = Pt(self.default_font_size)
    
    def _process_list_inline(
        self,
        list_element: Tag,
        paragraph,
        parent_formatting: Dict[str, bool],
        ordered: bool = False
    ) -> None:
        """Procesa listas HTML como texto con viñetas/números."""
        try:
            items = list_element.find_all('li', recursive=False)
            if not items:
                return
            
            if paragraph.text.strip():
                paragraph.add_run().add_break()
            
            for index, item in enumerate(items, start=1):
                if index > 1:
                    paragraph.add_run().add_break()
                
                prefix = f"{index}. " if ordered else "• "
                prefix_run = paragraph.add_run(prefix)
                if self.default_font_size:
                    prefix_run.font.size = Pt(self.default_font_size)
                
                self._process_list_item_content(item, paragraph, parent_formatting)
            
            self._needs_line_break = True
            
        except Exception as e:
            logger.error(f"Error procesando lista: {e}", exc_info=True)
    
    def _process_list_item_content(
        self,
        item_element: Tag,
        paragraph,
        parent_formatting: Dict[str, bool]
    ) -> None:
        """Procesa el contenido de un elemento <li>."""
        for child in item_element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    self._add_formatted_text(paragraph, text, parent_formatting)
            
            elif isinstance(child, Tag):
                tag_name = child.name.lower()
                new_formatting = parent_formatting.copy()
                
                if tag_name in ['strong', 'b']:
                    new_formatting['bold'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name in ['em', 'i']:
                    new_formatting['italic'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name == 'u':
                    new_formatting['underline'] = True
                    self._process_list_item_content(child, paragraph, new_formatting)
                
                elif tag_name == 'br':
                    paragraph.add_run().add_break()
                
                elif tag_name == 'p':
                    self._process_list_item_content(child, paragraph, parent_formatting)
                
                elif tag_name in ['span', 'a', 'div']:
                    self._process_list_item_content(child, paragraph, parent_formatting)
                
                elif tag_name in ['ul', 'ol']:
                    paragraph.add_run().add_break()
                    nested_ordered = (tag_name == 'ol')
                    nested_items = child.find_all('li', recursive=False)
                    for ni, nested_item in enumerate(nested_items, start=1):
                        if ni > 1:
                            paragraph.add_run().add_break()
                        nested_prefix = f"  {ni}. " if nested_ordered else "  ○ "
                        nested_run = paragraph.add_run(nested_prefix)
                        if self.default_font_size:
                            nested_run.font.size = Pt(self.default_font_size)
                        self._process_list_item_content(nested_item, paragraph, parent_formatting)
                
                else:
                    text = child.get_text(strip=True)
                    if text:
                        self._add_formatted_text(paragraph, text, parent_formatting)
    
    def _queue_table(self, table_element: Tag) -> None:
        """Encola una tabla para insertar después del párrafo."""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            table_data = {'rows': []}
            
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_data = []
                
                for cell in cells:
                    cell_text = cell.get_text(strip=True)
                    is_header = cell.name == 'th'
                    row_data.append({
                        'text': cell_text,
                        'bold': is_header
                    })
                
                if row_data:
                    table_data['rows'].append(row_data)
            
            if table_data['rows']:
                self._pending_tables.append(table_data)
                
        except Exception as e:
            logger.error(f"Error encolando tabla: {e}", exc_info=True)


# ============================================================================
# FUNCIONES AUXILIARES
# ============================================================================

def convert_html_to_word_inline(html_content: str, paragraph, document=None) -> None:
    """Función helper para convertir HTML a Word."""
    converter = HTMLToWordConverter()
    converter.convert_and_insert(html_content, paragraph, clear_paragraph=False, document=document)


def strip_html_tags(html_content: str) -> str:
    """Elimina tags HTML y devuelve solo el texto."""
    if not html_content:
        return ''
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text()
    except Exception:
        return html_content
